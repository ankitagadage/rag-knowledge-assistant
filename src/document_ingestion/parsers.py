"""
Format-specific parsers for the Document Ingestion Layer (SYSTEM_DESIGN.md
Section 3.1). Each parser takes a file path and returns a ParsedDocument
with plain text — the common shape everything downstream (chunking,
Section 3.1's next task) consumes, regardless of source format.

All parser-library-specific exceptions are caught and re-raised as
DocumentParseError, so callers only ever need to handle one exception type.
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Callable

import docx
import markdown as markdown_lib
import pypdf
from bs4 import BeautifulSoup
from docx.opc.exceptions import PackageNotFoundError
from pypdf.errors import PdfReadError

# Matches file_upload.allowed_types in SYSTEM_DESIGN.md Section 7.
SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".md"}


class DocumentParseError(Exception):
    """Raised for any failure while extracting text from a document."""


class UnsupportedDocumentType(DocumentParseError):
    """Raised when the file extension isn't one of SUPPORTED_EXTENSIONS."""


@dataclass
class ParsedDocument:
    text: str
    doc_type: str  # 'pdf' | 'txt' | 'docx' | 'md'
    page_count: int | None = None  # only meaningful for PDFs


def parse_pdf(path: Path) -> ParsedDocument:
    try:
        reader = pypdf.PdfReader(str(path))
    except FileNotFoundError:
        raise DocumentParseError(f"File not found: {path}")
    except PdfReadError as e:
        raise DocumentParseError(f"Corrupt or unreadable PDF: {path} ({e})")

    if reader.is_encrypted:
        raise DocumentParseError(f"PDF is password-protected: {path}")

    pages_text = [page.extract_text() or "" for page in reader.pages]
    text = "\n\n".join(pages_text).strip()

    if not text:
        raise DocumentParseError(
            f"No extractable text in PDF: {path} "
            f"(likely a scanned/image-only PDF — needs OCR, not supported yet)"
        )

    return ParsedDocument(text=text, doc_type="pdf", page_count=len(reader.pages))


def parse_docx(path: Path) -> ParsedDocument:
    try:
        document = docx.Document(str(path))
    except PackageNotFoundError:
        raise DocumentParseError(f"Corrupt or unreadable DOCX: {path}")
    except FileNotFoundError:
        raise DocumentParseError(f"File not found: {path}")

    text = "\n".join(p.text for p in document.paragraphs).strip()

    if not text:
        raise DocumentParseError(f"No extractable text in DOCX: {path}")

    return ParsedDocument(text=text, doc_type="docx")


def parse_txt(path: Path) -> ParsedDocument:
    try:
        text = path.read_text(encoding="utf-8").strip()
    except FileNotFoundError:
        raise DocumentParseError(f"File not found: {path}")
    except UnicodeDecodeError as e:
        raise DocumentParseError(f"Could not decode {path} as UTF-8 text: {e}")

    if not text:
        raise DocumentParseError(f"File is empty: {path}")

    return ParsedDocument(text=text, doc_type="txt")


def parse_markdown(path: Path) -> ParsedDocument:
    try:
        raw = path.read_text(encoding="utf-8").strip()
    except FileNotFoundError:
        raise DocumentParseError(f"File not found: {path}")
    except UnicodeDecodeError as e:
        raise DocumentParseError(f"Could not decode {path} as UTF-8 text: {e}")

    if not raw:
        raise DocumentParseError(f"File is empty: {path}")

    # Convert markdown -> HTML -> plain text, so chunk/embedding text isn't
    # cluttered with "##", "**", "[link](url)" syntax.
    html = markdown_lib.markdown(raw)
    text = BeautifulSoup(html, "html.parser").get_text().strip()

    return ParsedDocument(text=text, doc_type="md")


_PARSERS: dict[str, Callable[[Path], ParsedDocument]] = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".txt": parse_txt,
    ".md": parse_markdown,
}


def parse_document(path: str | Path) -> ParsedDocument:
    """Dispatches to the right parser based on file extension."""
    path = Path(path)
    extension = path.suffix.lower()

    parser = _PARSERS.get(extension)
    if parser is None:
        raise UnsupportedDocumentType(
            f"Unsupported file type '{extension}' for {path}. "
            f"Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    return parser(path)
